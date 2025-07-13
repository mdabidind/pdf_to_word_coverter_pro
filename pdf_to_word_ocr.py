#!/usr/bin/env python3

"""
PDF to Word OCR Converter

This script converts PDF files to editable Word documents using OCR (Optical Character Recognition).
It extracts text from each page of the PDF, converts the pages to images, performs OCR on the images,
and creates a Word document with the extracted text.

Requirements:
    - Python 3.6+
    - Tesseract OCR
    - Poppler (for pdf2image)
    - python-docx
    - pytesseract
    - pdf2image

Usage:
    python pdf_to_word_ocr.py input.pdf [-o output.docx] [-l lang] [-d dpi]

Options:
    -o, --output     Output Word document path (default: same as input with .docx extension)
    -l, --lang       OCR language (default: eng)
    -d, --dpi        DPI for image conversion (default: 300)
"""

import os
import sys
import argparse
import tempfile
from pathlib import Path
import logging
import platform

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger('pdf_to_word_ocr')

# Import required packages with error handling
try:
    from pdf2image import convert_from_path
except ImportError:
    logger.error("pdf2image package is not installed. Please install it using: pip install pdf2image")
    sys.exit(1)

try:
    import pytesseract
except ImportError:
    logger.error("pytesseract package is not installed. Please install it using: pip install pytesseract")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    logger.error("python-docx package is not installed. Please install it using: pip install python-docx")
    sys.exit(1)

try:
    from PIL import Image
except ImportError:
    logger.error("Pillow package is not installed. Please install it using: pip install Pillow")
    sys.exit(1)

def setup_tesseract():
    """Check if Tesseract is properly configured and set it up if needed"""
    try:
        # Try to get tesseract version
        version = pytesseract.get_tesseract_version()
        logger.info(f"Tesseract version: {version}")
        return True
    except Exception as e:
        logger.error(f"Tesseract not properly configured: {e}")
        
        # Try to automatically set up Tesseract path based on OS
        system = platform.system()
        if system == 'Windows':
            # Common installation paths on Windows
            common_paths = [
                r'C:\Program Files\Tesseract-OCR\tesseract.exe',
                r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
                r'C:\Tesseract-OCR\tesseract.exe'
            ]
            
            for path in common_paths:
                if os.path.exists(path):
                    pytesseract.pytesseract.tesseract_cmd = path
                    logger.info(f"Automatically set Tesseract path to: {path}")
                    return True
            
            logger.error("Could not find Tesseract installation.")
            logger.info("Please install Tesseract from: https://github.com/UB-Mannheim/tesseract/wiki")
            logger.info("Then set the path manually using: pytesseract.pytesseract.tesseract_cmd = r'path/to/tesseract.exe'")
        elif system == 'Linux':
            logger.info("On Linux, install Tesseract using: sudo apt-get install tesseract-ocr")
        elif system == 'Darwin':  # macOS
            logger.info("On macOS, install Tesseract using: brew install tesseract")
        
        return False

def convert_pdf_to_images(pdf_path, dpi=300):
    """Convert PDF pages to images"""
    logger.info(f"Converting PDF to images: {pdf_path}")
    try:
        # Create a temporary directory for the images
        with tempfile.TemporaryDirectory() as temp_dir:
            # Convert PDF to images
            images = convert_from_path(pdf_path, dpi=dpi, output_folder=temp_dir)
            logger.info(f"Converted {len(images)} pages to images")
            return images
    except Exception as e:
        logger.error(f"Error converting PDF to images: {e}")
        if "poppler" in str(e).lower():
            system = platform.system()
            if system == 'Windows':
                logger.info("On Windows, install poppler from: https://github.com/oschwartz10612/poppler-windows/releases")
                logger.info("Then add the bin directory to your PATH")
            elif system == 'Linux':
                logger.info("On Linux, install poppler using: sudo apt-get install poppler-utils")
            elif system == 'Darwin':  # macOS
                logger.info("On macOS, install poppler using: brew install poppler")
        return None

def perform_ocr(images, lang='eng'):
    """Perform OCR on the images"""
    logger.info(f"Performing OCR on {len(images)} images with language: {lang}")
    text_results = []
    
    for i, image in enumerate(images):
        logger.info(f"Processing page {i+1}/{len(images)}")
        try:
            text = pytesseract.image_to_string(image, lang=lang)
            text_results.append(text)
        except Exception as e:
            logger.error(f"Error performing OCR on page {i+1}: {e}")
            text_results.append("")
    
    return text_results

def create_word_document(text_results, output_path):
    """Create a Word document from the OCR results"""
    logger.info(f"Creating Word document: {output_path}")
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Add each page as a separate section
    for i, text in enumerate(text_results):
        if i > 0:
            # Add page break between pages
            doc.add_page_break()
        
        # Add heading for each page
        heading = doc.add_heading(f'Page {i+1}', level=2)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Process the text by paragraphs
        paragraphs = text.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                p = doc.add_paragraph(para_text.replace('\n', ' '))
    
    # Save the document
    try:
        # Ensure the output directory exists
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
            
        doc.save(output_path)
        logger.info(f"Word document saved successfully: {output_path}")
        return True
    except Exception as e:
        logger.error(f"Error saving Word document: {e}")
        return False

def convert_pdf_to_word(pdf_path, output_path=None, lang='eng', dpi=300):
    """Convert PDF to Word document using OCR"""
    # Validate input file
    pdf_path = Path(pdf_path)
    if not pdf_path.exists() or not pdf_path.is_file() or pdf_path.suffix.lower() != '.pdf':
        logger.error(f"Invalid PDF file: {pdf_path}")
        return False
    
    # Set default output path if not provided
    if output_path is None:
        output_path = pdf_path.with_suffix('.docx')
    else:
        output_path = Path(output_path)
        # Ensure output has .docx extension
        if output_path.suffix.lower() != '.docx':
            output_path = output_path.with_suffix('.docx')
    
    # Check if Tesseract is configured
    if not setup_tesseract():
        return False
    
    # Convert PDF to images
    images = convert_pdf_to_images(pdf_path, dpi)
    if not images:
        return False
    
    # Perform OCR on the images
    text_results = perform_ocr(images, lang)
    
    # Create Word document
    return create_word_document(text_results, output_path)

def main():
    # Parse command line arguments
    parser = argparse.ArgumentParser(description='Convert PDF to Word document using OCR')
    parser.add_argument('pdf_path', help='Path to the PDF file')
    parser.add_argument('-o', '--output', help='Output Word document path')
    parser.add_argument('-l', '--lang', default='eng', help='OCR language (default: eng)')
    parser.add_argument('-d', '--dpi', type=int, default=300, help='DPI for image conversion (default: 300)')
    args = parser.parse_args()
    
    # Convert PDF to Word
    success = convert_pdf_to_word(args.pdf_path, args.output, args.lang, args.dpi)
    
    if success:
        logger.info("Conversion completed successfully!")
        return 0
    else:
        logger.error("Conversion failed!")
        return 1

if __name__ == '__main__':
    sys.exit(main())